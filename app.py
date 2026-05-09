import hmac
import json
import os
import re
import tomllib

from collections.abc import Mapping

import streamlit as st
from streamlit.errors import StreamlitSecretNotFoundError
import base64
from io import BytesIO
import streamlit.components.v1 as components
import hashlib
from datetime import datetime, timezone
from fpdf import FPDF
from groq import Groq
from docx import Document
from pypdf import PdfReader
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity


DEFAULT_BRANDING: dict[str, object] = {
    "app_name": "Fluency",
    "tagline": "Capture expertise in a snap",
    "page_title": "Fluency",
    "page_icon": "🗣️",
    "logo_url": "",
    "logo_path": "",
    "primary_color": "#E1306C",
    "secondary_color": "#833AB4",
    "accent_color": "#FCAF45",
    "hide_powered_by": True,
}

_BRANDING_KEYS = frozenset(DEFAULT_BRANDING.keys())


def _normalize_secret_value(value: object) -> str | None:
    if value is None:
        return None
    s = str(value).strip()
    return s or None


def _secret_or_env(name: str) -> str | None:
    """Read a scalar secret from secrets.toml when present; fall back to the same-named env var."""
    try:
        if name in st.secrets:
            return _normalize_secret_value(st.secrets[name])
    except StreamlitSecretNotFoundError:
        pass
    return _normalize_secret_value(os.getenv(name))


def _optional_access_gate(brand: dict[str, object]) -> None:
    """If APP_ACCESS_PASSWORD is set, require it before rendering the rest of the app."""
    expected = _secret_or_env("APP_ACCESS_PASSWORD")
    if not expected:
        return
    if st.session_state.get("_access_granted"):
        return
    st.title(str(brand.get("app_name") or DEFAULT_BRANDING["app_name"]))
    st.caption("This deployment is password-protected.")
    with st.form("access_gate"):
        entered = st.text_input("Access password", type="password")
        if st.form_submit_button("Continue"):
            if hmac.compare_digest(
                entered.encode("utf-8"),
                expected.encode("utf-8"),
            ):
                st.session_state._access_granted = True
                st.rerun()
            else:
                st.error("Incorrect password.")
    st.stop()


def _coerce_bool(value: object, default: bool = False) -> bool:
    if isinstance(value, bool):
        return value
    s = str(value).strip().lower()
    if not s:
        return default
    return s in ("1", "true", "yes", "on")


def _safe_hex_color(value: object, default: str) -> str:
    if value is None:
        return default
    s = str(value).strip()
    if re.fullmatch(r"#([0-9A-Fa-f]{3}|[0-9A-Fa-f]{6})", s):
        return s
    return default


def _hex_to_rgb_tuple(hex_color: str) -> tuple[int, int, int]:
    h = hex_color.lstrip("#")
    if len(h) == 3:
        h = "".join(c * 2 for c in h)
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def _rgba(hex_color: str, alpha: float) -> str:
    r, g, b = _hex_to_rgb_tuple(hex_color)
    return f"rgba({r}, {g}, {b}, {alpha})"


def _branding_from_secrets_file() -> dict[str, object]:
    path = os.path.join(".streamlit", "secrets.toml")
    if not os.path.isfile(path):
        return {}
    try:
        with open(path, "rb") as f:
            data = tomllib.load(f)
    except (OSError, tomllib.TOMLDecodeError):
        return {}
    b = data.get("branding")
    return dict(b) if isinstance(b, dict) else {}


def _branding_from_streamlit_secrets() -> dict[str, object]:
    try:
        if "branding" not in st.secrets:
            return {}
        sec = st.secrets["branding"]
    except StreamlitSecretNotFoundError:
        return {}
    if not isinstance(sec, Mapping):
        return {}
    raw = dict(sec)
    return {str(k): v for k, v in raw.items() if str(k) in _BRANDING_KEYS}


def _branding_env_overrides() -> dict[str, object]:
    out: dict[str, object] = {}
    mapping = {
        "BRAND_APP_NAME": "app_name",
        "BRAND_TAGLINE": "tagline",
        "BRAND_PAGE_TITLE": "page_title",
        "BRAND_PAGE_ICON": "page_icon",
        "BRAND_LOGO_URL": "logo_url",
        "BRAND_LOGO_PATH": "logo_path",
        "BRAND_PRIMARY_COLOR": "primary_color",
        "BRAND_SECONDARY_COLOR": "secondary_color",
        "BRAND_ACCENT_COLOR": "accent_color",
        "BRAND_HIDE_POWERED_BY": "hide_powered_by",
    }
    for env_key, brand_key in mapping.items():
        val = os.getenv(env_key)
        if val is None or val.strip() == "":
            continue
        if brand_key == "hide_powered_by":
            out[brand_key] = _coerce_bool(val)
        else:
            out[brand_key] = val.strip()
    return out


def get_initial_branding() -> dict[str, object]:
    """Branding safe to compute before the first Streamlit command.

    Used for st.set_page_config so the favicon/tab icon can be customized.
    """
    merged: dict[str, object] = dict(DEFAULT_BRANDING)
    merged.update({k: v for k, v in _branding_from_secrets_file().items() if k in _BRANDING_KEYS})
    merged.update(_branding_env_overrides())
    merged["page_icon"] = str(merged.get("page_icon") or DEFAULT_BRANDING["page_icon"])
    merged["page_title"] = str(merged.get("page_title") or DEFAULT_BRANDING["page_title"])
    return merged


_initial_brand = get_initial_branding()
st.set_page_config(
    page_title=str(_initial_brand.get("page_title")),
    page_icon=str(_initial_brand.get("page_icon")),
    layout="wide",
)


def get_branding() -> dict[str, object]:
    merged: dict[str, object] = dict(DEFAULT_BRANDING)
    merged.update({k: v for k, v in _branding_from_secrets_file().items() if k in _BRANDING_KEYS})
    merged.update({k: v for k, v in _branding_from_streamlit_secrets().items() if k in _BRANDING_KEYS})
    merged.update(_branding_env_overrides())

    merged["primary_color"] = _safe_hex_color(
        merged.get("primary_color"), str(DEFAULT_BRANDING["primary_color"])
    )
    merged["secondary_color"] = _safe_hex_color(
        merged.get("secondary_color"), str(DEFAULT_BRANDING["secondary_color"])
    )
    merged["accent_color"] = _safe_hex_color(
        merged.get("accent_color"), str(DEFAULT_BRANDING["accent_color"])
    )
    merged["hide_powered_by"] = _coerce_bool(merged.get("hide_powered_by"), False)
    merged["page_icon"] = str(merged.get("page_icon") or DEFAULT_BRANDING["page_icon"])
    return merged


def _sync_browser_tab_title(title: str) -> None:
    safe = json.dumps(title or "SOP Generator")
    components.html(
        f"<script>try{{parent.document.title = {safe};}}catch(e){{}}</script>",
        height=0,
        width=0,
    )


def build_branding_css(brand: dict[str, object]) -> str:
    pr = str(brand.get("primary_color"))
    sec = str(brand.get("secondary_color"))
    ac = str(brand.get("accent_color"))

    shadow = _rgba(sec, 0.20)
    shadow_h = _rgba(sec, 0.26)
    shadow_a = _rgba(sec, 0.18)

    return f"""
<style>
/* ---- White-label UI (brand colors) ---- */
.stApp {{
  /* Dark, muted base so the app doesn't feel overly bright */
  background: radial-gradient(900px 500px at 10% 10%, {_rgba(pr, 0.10)}, transparent 60%),
              radial-gradient(800px 520px at 90% 20%, {_rgba(sec, 0.10)}, transparent 55%),
              radial-gradient(900px 600px at 50% 90%, {_rgba(ac, 0.08)}, transparent 60%),
              linear-gradient(180deg, #0B0F1A 0%, #0A0D14 100%);
  background-size: 120% 120%;
  animation: bgShift 14s ease-in-out infinite;
}}

@keyframes bgShift {{
  0%   {{ background-position: 0% 0%; }}
  50%  {{ background-position: 100% 40%; }}
  100% {{ background-position: 0% 0%; }}
}}

section[data-testid="stSidebar"] > div {{
  background: rgba(18, 24, 39, 0.72);
  backdrop-filter: blur(10px);
  border-right: 1px solid rgba(255, 255, 255, 0.08);
}}

div.block-container {{
  padding-top: 1.25rem;
  padding-bottom: 2.5rem;
}}

div.stButton > button {{
  border: 0;
  border-radius: 14px;
  padding: 0.65rem 1rem;
  background: linear-gradient(135deg, {pr} 0%, {sec} 55%, {ac} 100%);
  color: white !important;
  box-shadow: 0 10px 24px {shadow};
  transition: transform 120ms ease, box-shadow 120ms ease, filter 120ms ease;
}}
div.stButton > button:hover {{
  transform: translateY(-1px);
  box-shadow: 0 14px 28px {shadow_h};
  filter: saturate(1.05);
}}
div.stButton > button:active {{
  transform: translateY(0px) scale(0.99);
  box-shadow: 0 8px 18px {shadow_a};
}}

div[data-baseweb="input"] input,
div[data-baseweb="textarea"] textarea,
div[data-baseweb="select"] > div {{
  border-radius: 14px !important;
}}

div[data-testid="stExpander"] {{
  border-radius: 16px;
  border: 1px solid rgba(255,255,255,0.10);
  background: rgba(18, 24, 39, 0.55);
  backdrop-filter: blur(10px);
}}

h1, h2, h3 {{
  letter-spacing: -0.02em;
}}
</style>
"""


SVG_CODE = """
<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 240 80" width="240" height="80">
  <rect x="0" y="0" width="240" height="80" fill="#ffffff" rx="12" ry="12"/>
  <g transform="translate(20,40)">
    <circle cx="0" cy="0" r="24" fill="#0A74DA"/>
    <circle cx="8" cy="-6" r="12" fill="#ffffff"/>
  </g>
  <text x="70" y="48" font-family="Arial" font-size="36" font-weight="600" fill="#222222">SOP</text>
  <text x="70" y="70" font-family="Arial" font-size="14" fill="#555555">AI Generator</text>
</svg>
""".strip()


def render_svg_data_uri(svg: str) -> str:
    b64 = base64.b64encode(svg.encode("utf-8")).decode("utf-8")
    return f"data:image/svg+xml;base64,{b64}"


def resolve_brand_logo_url(brand: dict[str, object]) -> str:
    url = str(brand.get("logo_url") or "").strip()
    if url:
        return url
    path = str(brand.get("logo_path") or "").strip()
    if path:
        p = path if os.path.isabs(path) else os.path.join(os.getcwd(), path)
        if os.path.isfile(p):
            with open(p, "rb") as f:
                raw = f.read()
            ext = os.path.splitext(p)[1].lower()
            mime = {
                ".png": "image/png",
                ".jpg": "image/jpeg",
                ".jpeg": "image/jpeg",
                ".webp": "image/webp",
                ".svg": "image/svg+xml",
            }.get(ext, "application/octet-stream")
            b64 = base64.b64encode(raw).decode("ascii")
            return f"data:{mime};base64,{b64}"
    return render_svg_data_uri(SVG_CODE)


def header_tagline(brand: dict[str, object]) -> str | None:
    tag = str(brand.get("tagline") or "").strip()
    if _coerce_bool(brand.get("hide_powered_by"), False):
        return tag or None
    if tag:
        return tag
    return "Powered by Groq"


_brand = get_branding()
_tab_title = str(_brand.get("page_title") or _brand.get("app_name") or DEFAULT_BRANDING["page_title"])
_sync_browser_tab_title(_tab_title)
st.markdown(build_branding_css(_brand), unsafe_allow_html=True)

_optional_access_gate(_brand)

def create_pdf_bytes(text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    
    # Clean the text
    clean_text = (text or "").encode("latin-1", "ignore").decode("latin-1")
    pdf.multi_cell(0, 10, txt=clean_text)
    
    out = pdf.output(dest="S")
    # fpdf (PyFPDF) may return `str`, while fpdf2 often returns `bytes`/`bytearray`.
    if isinstance(out, str):
        return out.encode("latin-1")
    return bytes(out)


def create_docx_bytes(title: str, text: str) -> bytes:
    doc = Document()
    if title.strip():
        doc.add_heading(title.strip(), level=1)

    for raw_line in (text or "").splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        else:
            doc.add_paragraph(line)

    buff = BytesIO()
    doc.save(buff)
    return buff.getvalue()


def sop_fingerprint(text: str) -> str:
    return hashlib.sha256((text or "").encode("utf-8")).hexdigest()


def _chunk_text(text: str, *, chunk_chars: int = 900, overlap_chars: int = 150) -> list[str]:
    t = (text or "").strip()
    if not t:
        return []
    chunks: list[str] = []
    i = 0
    while i < len(t):
        end = min(len(t), i + chunk_chars)
        chunk = t[i:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(t):
            break
        i = max(0, end - overlap_chars)
    return chunks


@st.cache_data(show_spinner=False, ttl=3600, max_entries=128)
def extract_pdf_text_cached(*, file_sha256: str, pdf_bytes: bytes) -> str:
    # file_sha256 is part of the cache key; pdf_bytes is the payload.
    reader = PdfReader(BytesIO(pdf_bytes))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")
    return "\n".join(parts).strip()


@st.cache_data(show_spinner=False, ttl=3600, max_entries=64)
def build_rag_index_cached(*, corpus_sha256: str, chunks: list[str]) -> dict:
    # Store vectorizer vocabulary + matrix; st.cache_data will pickle it.
    vectorizer = TfidfVectorizer(stop_words="english", max_features=5000)
    matrix = vectorizer.fit_transform(chunks) if chunks else None
    return {"vectorizer": vectorizer, "matrix": matrix}


def retrieve_company_snippets(
    *,
    query: str,
    docs: list[dict],
    top_k: int = 6,
) -> list[dict]:
    # docs: [{name, chunks:[...]}]
    all_chunks: list[str] = []
    meta: list[dict] = []
    for d in docs:
        name = d.get("name", "manual.pdf")
        for idx, ch in enumerate(d.get("chunks", []) or []):
            all_chunks.append(ch)
            meta.append({"doc": name, "chunk_index": idx, "text": ch})

    if not all_chunks or not query.strip():
        return []

    corpus_sha = hashlib.sha256(("\n".join(all_chunks)).encode("utf-8")).hexdigest()
    index = build_rag_index_cached(corpus_sha256=corpus_sha, chunks=all_chunks)
    vectorizer: TfidfVectorizer = index["vectorizer"]
    matrix = index["matrix"]
    if matrix is None:
        return []

    qv = vectorizer.transform([query])
    sims = cosine_similarity(qv, matrix)[0]
    ranked = sorted(range(len(sims)), key=lambda i: sims[i], reverse=True)[:top_k]

    results: list[dict] = []
    for i in ranked:
        if sims[i] <= 0:
            continue
        m = meta[i]
        results.append(
            {
                "doc": m["doc"],
                "chunk_index": m["chunk_index"],
                "score": float(sims[i]),
                "text": m["text"],
            }
        )
    return results



def get_groq_api_key() -> str | None:
    """Resolve Groq API key from Streamlit secrets.toml (or env injected by Streamlit), then bare env."""
    try:
        if "GROQ_API_KEY" in st.secrets:
            k = _normalize_secret_value(st.secrets["GROQ_API_KEY"])
            if k:
                return k
        if "groq" in st.secrets:
            section = st.secrets["groq"]
            if isinstance(section, Mapping) and "api_key" in section:
                k = _normalize_secret_value(section["api_key"])
                if k:
                    return k
    except StreamlitSecretNotFoundError:
        pass

    return _normalize_secret_value(os.getenv("GROQ_API_KEY"))


COMPANY_PROFILE_PATH = os.path.join(".streamlit", "company_profile.json")
FEEDBACK_PATH = os.path.join(".streamlit", "feedback.jsonl")
HISTORY_PATH = os.path.join(".streamlit", "history.json")


def load_company_profile() -> dict:
    try:
        if not os.path.exists(COMPANY_PROFILE_PATH):
            return {}
        with open(COMPANY_PROFILE_PATH, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def save_company_profile(profile: dict) -> None:
    os.makedirs(os.path.dirname(COMPANY_PROFILE_PATH), exist_ok=True)
    with open(COMPANY_PROFILE_PATH, "w", encoding="utf-8") as f:
        json.dump(profile, f, ensure_ascii=False, indent=2)


def show_busy_error() -> None:
    st.error("The system is busy. Please try again in a few seconds.")


def append_feedback(entry: dict) -> None:
    os.makedirs(os.path.dirname(FEEDBACK_PATH), exist_ok=True)
    with open(FEEDBACK_PATH, "a", encoding="utf-8") as f:
        f.write(json.dumps(entry, ensure_ascii=False) + "\n")


def load_recent_feedback(limit: int = 50) -> list[dict]:
    try:
        if not os.path.exists(FEEDBACK_PATH):
            return []
        with open(FEEDBACK_PATH, "r", encoding="utf-8") as f:
            lines = f.readlines()
        out: list[dict] = []
        for line in lines[-limit:]:
            line = line.strip()
            if not line:
                continue
            try:
                obj = json.loads(line)
                if isinstance(obj, dict):
                    out.append(obj)
            except Exception:
                continue
        return out
    except Exception:
        return []


def load_history() -> list[dict]:
    try:
        if not os.path.exists(HISTORY_PATH):
            return []
        with open(HISTORY_PATH, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, list) else []
    except Exception:
        return []


def save_history(items: list[dict]) -> None:
    os.makedirs(os.path.dirname(HISTORY_PATH), exist_ok=True)
    with open(HISTORY_PATH, "w", encoding="utf-8") as f:
        json.dump(items, f, ensure_ascii=False, indent=2)


def add_to_history(entry: dict, *, limit: int = 5) -> None:
    items = load_history()
    items.insert(0, entry)
    # de-dupe by sop_sha256 if present
    seen: set[str] = set()
    deduped: list[dict] = []
    for it in items:
        key = str(it.get("sop_sha256") or "")
        if key and key in seen:
            continue
        if key:
            seen.add(key)
        deduped.append(it)
    save_history(deduped[:limit])


TEMPLATE_GUIDANCE: dict[str, str] = {
    "IT SOP": """
Focus on technical accuracy, security, and repeatability.
Include: Preconditions/requirements, access/permissions, tools/systems involved,
rollback plan, troubleshooting, validation steps, logging/monitoring, and SLAs/owners.
Add a short 'Change management' section (impact, approvals, maintenance window).
""".strip(),
    "HR SOP": """
Focus on compliance, privacy, fairness, and a clear human workflow.
Include: Trigger events, required forms/documents, approvals, timelines/SLAs,
confidentiality/data handling, escalation paths, templates/communications, and record retention.
Add a short 'Candidate/employee communication' checklist.
""".strip(),
    "Warehouse SOP": """
Focus on safety, efficiency, and physical process clarity.
Include: PPE/safety requirements, equipment/tools, location/bin labeling, scanning steps,
quality checks, exception handling (damages/shortages), and end-of-shift reconciliation.
Add a short 'Safety checks' section and 'Common errors to avoid'.
""".strip(),
    "Restaurant SOP": """
Focus on food safety, service consistency, and speed.
Include: food safety controls (temps, cross-contamination), prep/line setup,
service steps, cleaning schedules, allergen handling, customer escalation, and close-down tasks.
Add checklists for opening/shift/closing and 'Quality standards' (taste, plating, timing).
""".strip(),
}


def build_prompt_for_template(
    template_name: str,
    topic: str,
    notes: str,
    *,
    audience: str,
    tools_used: str,
    compliance_standard: str,
    strictness: str,
    tone: str,
    include_definitions: bool,
    include_safety_compliance: bool,
    include_records: bool,
    include_checklist: bool,
) -> str:
    template_guidance = TEMPLATE_GUIDANCE.get(template_name, "")
    feedback_items = load_recent_feedback(limit=60)
    recent_down_reasons = [
        (it.get("reason") or "").strip()
        for it in feedback_items
        if it.get("rating") == "down" and it.get("template_name") == template_name
    ]
    recent_down_reasons = [r for r in recent_down_reasons if r][:3]
    feedback_avoid = ""
    if recent_down_reasons:
        bullets = "\n".join([f"- {r}" for r in recent_down_reasons])
        feedback_avoid = f"""
Common issues to avoid (from user feedback on previous SOPs):
{bullets}
""".strip()

    strictness_instructions = (
        "Strictness: STRICT. Use a formal, policy-like tone. Use short, unambiguous steps. "
        "Avoid fluff. Prefer MUST/SHALL where appropriate. Include clear acceptance/verification criteria."
        if strictness == "Strict"
        else "Strictness: DETAILED. Be thorough and explanatory while staying professional. "
        "Include tips, examples, and clarifying notes where helpful."
    )

    section_lines = [
        "- Purpose",
        "- Scope",
        "- Roles & responsibilities",
        "- Procedure (numbered)",
        "- Exceptions / edge cases",
    ]
    if include_definitions:
        section_lines.append("- Definitions (only if needed)")
    if include_records:
        section_lines.append("- Records / documentation")
    if include_safety_compliance:
        section_lines.append("- Safety / compliance (if relevant)")
    if include_checklist:
        section_lines.append("- Checklist (short, at the end)")

    sections_text = "\n".join(section_lines)

    notes_based_section = ""
    if len((notes or "").strip()) >= 1200:
        notes_based_section = """
Add a short section near the top titled exactly: "Based on notes:"
- 5–10 bullet points capturing the most important concrete facts from the notes (tools, roles, constraints, risks, timelines).
- Each bullet should reference the notes by quoting a short phrase in double-quotes OR by citing a specific detail (names/titles are ok; avoid secrets).
- Do NOT invent details that are not present in the notes. If something is missing, say "Not specified in notes".
""".strip()
    company_rules = st.session_state.get("company_rules_context", "").strip()
    company_rules_block = ""
    if company_rules:
        company_rules_block = f"""
Company rules (from uploaded manuals) — FOLLOW THESE STRICTLY:
{company_rules}
If any manual rule conflicts with the user's notes, call out the conflict and choose the safer/compliant path.
""".strip()

    return f"""
Write a clear, professional Standard Operating Procedure (SOP) for the topic below.
Use crisp headings and numbered steps. Keep it practical and immediately actionable.

Target audience: {audience}
Tools/systems used: {tools_used or "Not specified"}
Compliance standard(s): {compliance_standard or "Not specified"}
Tone: {tone}
{strictness_instructions}

{feedback_avoid}

{company_rules_block}

{notes_based_section}

Required sections (include ONLY these; omit all others):
{sections_text}

Template-specific guidance:
{template_guidance}

Topic: {topic}
Notes / raw input (may be messy): {notes}
""".strip()


@st.cache_data(show_spinner=False, ttl=3600, max_entries=256)
def generate_sop_cached(
    *,
    api_key: str,
    model: str,
    temperature: float,
    prompt: str,
) -> str:
    client = Groq(api_key=api_key)
    completion = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "You are a professional technical writer."},
            {"role": "user", "content": prompt},
        ],
        temperature=float(temperature),
    )
    return completion.choices[0].message.content or ""


@st.cache_data(show_spinner=False, ttl=3600, max_entries=256)
def review_and_fix_sop_cached(
    *,
    api_key: str,
    model: str,
    temperature: float,
    sop_text: str,
    strictness: str,
    tone: str,
    compliance_standard: str,
) -> str:
    client = Groq(api_key=api_key)
    prompt = f"""
You are reviewing an SOP for quality and completeness.

Goals:
- Find and fix gaps, unclear steps, missing roles/responsibilities, and missing records/documentation.
- Ensure steps are testable/verifyable and ordered logically.
- Ensure compliance language is appropriate for: {compliance_standard or "Not specified"} (if any).
- Keep the same overall intent, but rewrite as a corrected, improved SOP.
- If there is a "Based on notes:" section, keep it and correct it to match the SOP (do not add new facts).

Output rules:
- Return ONLY the revised SOP (no analysis, no bullet list of issues).
- Use the same tone: {tone}
- Use strictness: {strictness}

SOP to review:
{sop_text}
""".strip()

    completion = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "You are a meticulous SOP editor and auditor."},
            {"role": "user", "content": prompt},
        ],
        temperature=float(temperature),
    )
    return completion.choices[0].message.content or ""


@st.cache_data(show_spinner=False, ttl=3600, max_entries=256)
def compliance_audit_cached(
    *,
    api_key: str,
    model: str,
    temperature: float,
    sop_text: str,
    template_name: str,
    strictness: str,
    tone: str,
    compliance_standard: str,
    company_rules_context: str,
) -> str:
    client = Groq(api_key=api_key)
    prompt = f"""
You are a Compliance Auditor (AI Critic). Review the SOP below and critique it.

Context:
- Template: {template_name}
- Compliance standard(s): {compliance_standard or "Not specified"}
- Strictness: {strictness}
- Tone: {tone}

Company rules (if provided; treat as policy requirements):
{company_rules_context or "None"}

What to check:
- Compliance gaps or risky/unsafe steps
- Missing roles/responsibilities, approvals, evidence/records
- Unclear, untestable, or ambiguous steps
- Missing exceptions/edge cases and escalation paths
- Missing safety controls (PPE, cross-contamination, PHI/PII, access control, etc.) when relevant
- Conflicts with company rules (if any) and how to resolve them

Output format (use these headings exactly):
## Summary
## Findings (ranked)
- [High] ...
- [Medium] ...
- [Low] ...
## Missing items checklist
- ...
## Recommended fixes
1. ...

Return only the critique (no extra commentary).

SOP:
{sop_text}
""".strip()

    completion = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "You are a strict compliance auditor and SOP critic."},
            {"role": "user", "content": prompt},
        ],
        temperature=float(temperature),
    )
    return completion.choices[0].message.content or ""

@st.cache_data(show_spinner=False, ttl=3600, max_entries=256)
def generate_flowchart_mermaid_cached(
    *,
    api_key: str,
    model: str,
    temperature: float,
    sop_text: str,
) -> str:
    client = Groq(api_key=api_key)
    prompt = f"""
Create a Mermaid flowchart for the SOP below.

Rules:
- Output ONLY Mermaid code.
- Start with: flowchart TD
- Keep it readable: at most ~18 nodes.
- Use decision diamonds with labels like "Yes"/"No" paths when needed.
- Include start/end nodes.
- Do NOT include markdown fences.

SOP:
{sop_text}
""".strip()

    completion = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "You convert SOPs into clear flowcharts."},
            {"role": "user", "content": prompt},
        ],
        temperature=float(temperature),
    )
    return (completion.choices[0].message.content or "").strip()


def render_mermaid(mermaid_code: str, *, height_px: int = 700) -> None:
    code = (mermaid_code or "").strip()
    if not code:
        st.info("No flowchart to display.")
        return

    # Mermaid is rendered client-side via CDN.
    html = f"""
<div class="mermaid">
{code}
</div>
<script type="module">
  import mermaid from "https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.esm.min.mjs";
  mermaid.initialize({{
    startOnLoad: true,
    theme: "default",
    flowchart: {{ curve: "basis" }}
  }});
</script>
"""
    components.html(html, height=height_px, scrolling=True)


@st.cache_data(show_spinner=False, ttl=3600, max_entries=256)
def transcribe_audio_cached(
    *,
    api_key: str,
    model: str,
    file_name: str,
    file_sha256: str,
    audio_bytes: bytes,
    language: str,
) -> str:
    client = Groq(api_key=api_key)
    transcription = client.audio.transcriptions.create(
        file=(file_name, audio_bytes),
        model=model,
        response_format="json",
        language=language or None,
        temperature=0.0,
    )
    # Groq SDK returns an object with .text
    return (getattr(transcription, "text", None) or "").strip()


@st.cache_data(show_spinner=False, ttl=3600, max_entries=256)
def analyze_image_to_notes_cached(
    *,
    api_key: str,
    model: str,
    file_sha256: str,
    mime_type: str,
    image_b64: str,
) -> str:
    client = Groq(api_key=api_key)
    prompt = """
You are extracting actionable SOP notes from an image.

Return concise NOTES ONLY (no preamble), as bullet points grouped by:
- What is shown
- Key entities (tools/systems/roles)
- Steps / sequence (if implied)
- Requirements / constraints
- Risks / safety / compliance signals
- Any numbers, dates, thresholds, or checklists visible

If the image is a form/table/screenshot, capture the important fields and values.
Do not invent details; if unclear, say "Unclear in image".
""".strip()

    completion = client.chat.completions.create(
        model=model,
        messages=[
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:{mime_type};base64,{image_b64}"},
                    },
                ],
            }
        ],
        temperature=0.1,
    )
    return (completion.choices[0].message.content or "").strip()

logo_url = resolve_brand_logo_url(_brand)

with st.sidebar:
    # Load profile once per session, then use it as widget defaults.
    if "company_profile_loaded" not in st.session_state:
        profile = load_company_profile()
        st.session_state.company_profile_loaded = True
        st.session_state.profile_audience = str(profile.get("audience", "") or "")
        st.session_state.profile_tools_used = str(profile.get("tools_used", "") or "")
        st.session_state.profile_compliance = str(profile.get("compliance_standard", "") or "")
        st.session_state.profile_tone = str(profile.get("tone", "Professional") or "Professional")

    st.markdown("## Professional Edition")

    _has_custom_logo = bool(str(_brand.get("logo_url") or "").strip() or str(_brand.get("logo_path") or "").strip())
    if _has_custom_logo:
        st.image(logo_url, width=160)
    else:
        st.markdown(
            """
<div style="border: 1px dashed rgba(255,255,255,0.22); border-radius: 16px; padding: 18px; text-align: center; background: rgba(255,255,255,0.04);">
  <div style="font-weight: 700; letter-spacing: 0.04em; opacity: 0.9;">LOGO</div>
  <div style="margin-top: 6px; font-size: 12px; opacity: 0.7;">Upload / configure a logo</div>
</div>
""",
            unsafe_allow_html=True,
        )

    st.markdown(f"**{str(_brand.get('app_name') or DEFAULT_BRANDING['app_name'])}**")
    st.markdown("### How to use")
    st.info(
        "1. Enter a clear **Topic**.\n"
        "2. Paste **raw notes** or a transcript.\n"
        "3. Click **Generate SOP**.\n"
        "4. Download as PDF if needed."
    )

    st.markdown("### Company profile")
    st.text_input(
        "Default audience",
        key="profile_audience",
        placeholder="e.g., New hires, IT admins, Shift supervisors",
    )
    st.text_input(
        "Default tools used",
        key="profile_tools_used",
        placeholder="e.g., Okta, Jira, Google Workspace, Forklifts, POS system",
    )
    st.selectbox(
        "Default compliance",
        ["", "ISO 27001", "SOC 2", "HIPAA"],
        key="profile_compliance",
        index=0,
    )
    st.selectbox(
        "Default tone",
        ["Professional", "Friendly", "Policy-like", "Concise"],
        key="profile_tone",
        index=0,
    )

    col_p1, col_p2 = st.columns(2)
    with col_p1:
        if st.button("Save profile"):
            save_company_profile(
                {
                    "audience": st.session_state.profile_audience,
                    "tools_used": st.session_state.profile_tools_used,
                    "compliance_standard": st.session_state.profile_compliance,
                    "tone": st.session_state.profile_tone,
                }
            )
            st.success("Saved.")
    with col_p2:
        if st.button("Reset profile"):
            st.session_state.profile_audience = ""
            st.session_state.profile_tools_used = ""
            st.session_state.profile_compliance = ""
            st.session_state.profile_tone = "Professional"
            save_company_profile(
                {
                    "audience": "",
                    "tools_used": "",
                    "compliance_standard": "",
                    "tone": "Professional",
                }
            )
            st.success("Reset.")

    st.markdown("### Company Brain (RAG Lite)")
    manuals = st.file_uploader(
        "Upload PDF manuals (optional)",
        type=["pdf"],
        accept_multiple_files=True,
        help="Examples: Employee Handbook, Safety Guidelines, IT policy. These will be used as strict rules for SOPs.",
    )
    rag_top_k = st.slider("Manual snippets to use", 2, 10, 6, 1)

    if "company_manual_docs" not in st.session_state:
        st.session_state.company_manual_docs = []

    if manuals:
        docs: list[dict] = []
        for f in manuals:
            try:
                pdf_bytes = f.getvalue()
                sha = hashlib.sha256(pdf_bytes).hexdigest()
                text = extract_pdf_text_cached(file_sha256=sha, pdf_bytes=pdf_bytes)
                chunks = _chunk_text(text, chunk_chars=900, overlap_chars=150)
                docs.append({"name": f.name, "sha256": sha, "chunks": chunks})
            except Exception:
                continue
        st.session_state.company_manual_docs = docs
        st.write(f"Loaded manuals: {len(docs)}")
    else:
        st.caption("No manuals uploaded.")

    st.markdown("### History (last 5)")
    history_items = load_history()
    if history_items:
        options = []
        for i, it in enumerate(history_items):
            ts = (it.get("ts") or "")[:19].replace("T", " ")
            label = it.get("label") or it.get("template_name") or "SOP"
            options.append(f"{i+1}. {label} — {ts}")

        selected = st.selectbox("Saved SOPs", options, index=0)
        sel_idx = int(selected.split(".")[0]) - 1
        selected_item = history_items[sel_idx]

        if st.button("Load into editor"):
            st.session_state.current_sop_text = selected_item.get("sop_text", "") or ""
            st.session_state.last_inferred_topic = selected_item.get("label", "SOP") or "SOP"
            st.success("Loaded into editor.")

        col_h1, col_h2 = st.columns(2)
        with col_h1:
            if st.button("Delete selected"):
                remaining = [it for j, it in enumerate(history_items) if j != sel_idx]
                save_history(remaining[:5])
                st.success("Deleted.")
        with col_h2:
            if st.button("Clear history"):
                save_history([])
                st.success("Cleared.")
    else:
        st.caption("No saved SOPs yet.")

    st.markdown("### Settings")
    template_name = st.selectbox(
        "Template",
        ["IT SOP", "HR SOP", "Warehouse SOP", "Restaurant SOP"],
        index=0,
    )

    strictness = st.radio("Strictness", ["Strict", "Detailed"], index=1, horizontal=True)

    audience = st.text_input(
        "Audience (optional)",
        value=st.session_state.profile_audience,
        placeholder="e.g., New hires, IT admins, Shift supervisors",
    )
    tools_used = st.text_input(
        "Tools used (optional)",
        value=st.session_state.profile_tools_used,
        placeholder="e.g., Okta, Jira, Google Workspace, Forklifts, POS system",
    )
    compliance_standard = st.selectbox(
        "Compliance standard (optional)",
        ["None", "ISO 27001", "SOC 2", "HIPAA"],
        index=0,
    )
    compliance_standard = "" if compliance_standard == "None" else compliance_standard
    tone = st.selectbox(
        "Tone",
        ["Professional", "Friendly", "Policy-like", "Concise"],
        index=["Professional", "Friendly", "Policy-like", "Concise"].index(st.session_state.profile_tone)
        if st.session_state.profile_tone in ["Professional", "Friendly", "Policy-like", "Concise"]
        else 0,
    )

    st.markdown("### Outline controls")
    include_definitions = st.checkbox("Include Definitions section", value=True)
    include_safety_compliance = st.checkbox("Include Safety/Compliance section", value=True)
    include_records = st.checkbox("Include Records/Documentation section", value=True)
    include_checklist = st.checkbox("Include Checklist section", value=True)

    # Fixed temperature for stable outputs (removed "Creativity level" control)
    temperature = 0.35
    model="llama-3.1-8b-instant"

    if st.button("Clear cached results"):
        st.cache_data.clear()


header_left, header_right = st.columns([1, 6])
with header_left:
    st.image(logo_url, width=70)
with header_right:
    st.title(str(_brand.get("app_name") or DEFAULT_BRANDING["app_name"]))
    _tag = header_tagline(_brand)
    if _tag:
        st.caption(_tag)

if "notes" not in st.session_state:
    st.session_state.notes = ""

api_key = get_groq_api_key()
if not api_key:
    st.warning(
        "Set `GROQ_API_KEY` in `.streamlit/secrets.toml` (see `.streamlit/secrets.toml.example`) "
        "or as the environment variable `GROQ_API_KEY` to generate SOPs."
    )

with st.expander("Voice Mode (Audio-to-SOP)", expanded=False):
    st.caption("Upload an audio file, transcribe it, then generate the SOP from the transcript.")
    audio_file = st.file_uploader(
        "Upload audio",
        type=["wav", "mp3", "m4a", "aac", "flac", "ogg", "webm"],
        accept_multiple_files=False,
    )
    stt_model = st.selectbox(
        "Speech-to-text model",
        ["whisper-large-v3-turbo", "whisper-large-v3"],
        index=0,
    )
    stt_language = st.text_input("Language (optional, ISO-639-1)", value="", placeholder="e.g., en")

    if st.button("Transcribe audio", disabled=(not api_key or audio_file is None)):
        try:
            audio_bytes = audio_file.getvalue()
            file_sha = hashlib.sha256(audio_bytes).hexdigest()
            with st.spinner("Transcribing..."):
                transcript = transcribe_audio_cached(
                    api_key=api_key,
                    model=stt_model,
                    file_name=audio_file.name,
                    file_sha256=file_sha,
                    audio_bytes=audio_bytes,
                    language=stt_language.strip(),
                )
            if transcript:
                st.session_state.notes = transcript
                st.success("Transcription complete. The Notes box below was filled.")
            else:
                st.error("Transcription returned empty text.")
        except Exception:
            show_busy_error()
with st.expander("Vision (Image Analysis)", expanded=False):
    st.caption("Upload an image (photo/screenshot). We'll extract structured notes and fill the Notes box.")
    image_file = st.file_uploader(
        "Upload image",
        type=["png", "jpg", "jpeg", "webp"],
        accept_multiple_files=False,
    )
    vision_model = st.selectbox(
        "Vision model",
        ["meta-llama/llama-4-scout-17b-16e-instruct"],
        index=0,
    )

    if image_file is not None:
        st.image(image_file, caption=image_file.name, use_container_width=True)

    if st.button("Analyze image", disabled=(not api_key or image_file is None)):
        try:
            image_bytes = image_file.getvalue()
            file_sha = hashlib.sha256(image_bytes).hexdigest()
            mime_type = image_file.type or "image/png"
            image_b64 = base64.b64encode(image_bytes).decode("utf-8")

            with st.spinner("Analyzing image..."):
                extracted_notes = analyze_image_to_notes_cached(
                    api_key=api_key,
                    model=vision_model,
                    file_sha256=file_sha,
                    mime_type=mime_type,
                    image_b64=image_b64,
                )

            if extracted_notes:
                st.session_state.notes = extracted_notes
                st.success("Image analysis complete. The Notes box below was filled.")
            else:
                st.error("Image analysis returned empty text.")
        except Exception:
            show_busy_error()

notes = st.text_area(
    "Input notes / raw text",
    key="notes",
    height=220,
    placeholder="Paste your notes here (or use Voice Mode / Vision to generate notes).",
)

generate = st.button("Generate SOP", type="primary", disabled=not api_key)

if generate:
    if not notes.strip():
        st.error("Please paste your notes (or a transcript) first.")
    else:
        with st.spinner("Writing SOP..."):
            try:
                inferred_topic = f"{template_name} SOP"
                # Build company brain context for this generation (stored in session_state).
                company_docs = st.session_state.get("company_manual_docs", []) or []
                st.session_state.company_rules_context = ""
                if company_docs:
                    query = f"{template_name}\n{inferred_topic}\n{audience}\n{tools_used}\n{compliance_standard}\n{notes}"
                    snippets = retrieve_company_snippets(query=query, docs=company_docs, top_k=int(rag_top_k))
                    if snippets:
                        ctx_lines = []
                        for s in snippets:
                            ctx_lines.append(
                                f"- ({s['doc']} #chunk{s['chunk_index']}) {s['text']}"
                            )
                        st.session_state.company_rules_context = "\n".join(ctx_lines)

                prompt = build_prompt_for_template(
                    template_name,
                    inferred_topic,
                    notes,
                    audience=audience.strip() or "General staff",
                    tools_used=tools_used.strip(),
                    compliance_standard=compliance_standard.strip(),
                    strictness=strictness,
                    tone=tone,
                    include_definitions=include_definitions,
                    include_safety_compliance=include_safety_compliance,
                    include_records=include_records,
                    include_checklist=include_checklist,
                )
                sop_text = generate_sop_cached(
                    api_key=api_key,
                    model=model,
                    temperature=float(temperature),
                    prompt=prompt,
                )
                st.session_state.last_sop_text = sop_text
                st.session_state.last_inferred_topic = inferred_topic
                # Always set the "current" SOP so it persists across any subsequent button clicks.
                st.session_state.current_sop_text = sop_text

                add_to_history(
                    {
                        "ts": datetime.now(timezone.utc).isoformat(),
                        "label": inferred_topic,
                        "template_name": template_name,
                        "source": "generated",
                        "sop_text": sop_text,
                        "sop_sha256": sop_fingerprint(sop_text),
                    }
                )
            except Exception as e:
                show_busy_error()
                sop_text = ""

        if sop_text:
            st.subheader("Generated SOP")
            st.markdown(sop_text)

            with st.expander("Interactive Step Editor (edit inside the app)", expanded=False):
                st.caption("Edit the SOP here, then download the edited version.")
                edited = st.text_area(
                    "Edit SOP text",
                    value=st.session_state.current_sop_text,
                    height=320,
                    key=f"editor_gen_{sop_fingerprint(sop_text)}",
                )
                col_e1, col_e2 = st.columns(2)
                with col_e1:
                    if st.button("Save edits", key=f"save_gen_{sop_fingerprint(sop_text)}"):
                        st.session_state.current_sop_text = edited
                        st.success("Edits saved. Downloads will use the edited SOP.")
                with col_e2:
                    if st.button("Reset to generated", key=f"reset_gen_{sop_fingerprint(sop_text)}"):
                        st.session_state.current_sop_text = sop_text
                        st.success("Reset to the generated SOP.")

            sop_for_download = st.session_state.get("current_sop_text") or sop_text
            safe_name = "".join(c for c in inferred_topic.strip() if c.isalnum() or c in (" ", "-", "_")).strip() or "sop"
            try:
                pdf_bytes = create_pdf_bytes(sop_for_download)
                docx_bytes = create_docx_bytes(safe_name, sop_for_download)
            except Exception:
                pdf_bytes = b""
                docx_bytes = b""
                show_busy_error()

            col_a, col_b = st.columns(2)
            with col_a:
                if pdf_bytes:
                    st.download_button(
                        "Download PDF",
                        data=pdf_bytes,
                        file_name=f"{safe_name}.pdf",
                        mime="application/pdf",
                        key=f"dl_pdf_gen_{sop_fingerprint(sop_for_download)}",
                    )
            with col_b:
                if docx_bytes:
                    st.download_button(
                        "Download DOCX",
                        data=docx_bytes,
                        file_name=f"{safe_name}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"dl_docx_gen_{sop_fingerprint(sop_for_download)}",
                    )

            st.markdown("### Rate this SOP")
            rating = st.radio(
                "Was this SOP helpful?",
                ["👍 Thumbs Up", "👎 Thumbs Down"],
                horizontal=True,
                key=f"rating_{sop_fingerprint(sop_text)}",
            )
            reason = ""
            if rating.startswith("👎"):
                reason = st.text_area(
                    "What should be improved?",
                    placeholder="e.g., missing roles, unclear steps, wrong order, missing records, too long/short...",
                    key=f"reason_{sop_fingerprint(sop_text)}",
                )
            if st.button("Submit feedback", key=f"submit_{sop_fingerprint(sop_text)}"):
                append_feedback(
                    {
                        "ts": datetime.now(timezone.utc).isoformat(),
                        "rating": "up" if rating.startswith("👍") else "down",
                        "reason": reason.strip(),
                        "template_name": template_name,
                        "strictness": strictness,
                        "tone": tone,
                        "compliance_standard": compliance_standard or "",
                        "audience": (audience or "").strip(),
                        "tools_used": (tools_used or "").strip(),
                        "include_definitions": bool(include_definitions),
                        "include_safety_compliance": bool(include_safety_compliance),
                        "include_records": bool(include_records),
                        "include_checklist": bool(include_checklist),
                        "model": model,
                        "temperature": float(temperature),
                        "notes_chars": len((notes or "").strip()),
                        "sop_sha256": sop_fingerprint(sop_text),
                    }
                )
                st.success("Thanks — feedback saved.")


# --- Persistent SOP display (state management) ---
current_sop = (st.session_state.get("current_sop_text") or "").strip()
if current_sop:
    st.divider()
    st.subheader("Current SOP")
    st.markdown(current_sop)

    inferred_topic = str(st.session_state.get("last_inferred_topic") or "sop")
    safe_name = "".join(c for c in inferred_topic.strip() if c.isalnum() or c in (" ", "-", "_")).strip() or "sop"
    try:
        current_pdf = create_pdf_bytes(current_sop)
    except Exception:
        current_pdf = b""
        show_busy_error()

    if current_pdf:
        st.download_button(
            "Download PDF",
            data=current_pdf,
            file_name=f"{safe_name}.pdf",
            mime="application/pdf",
            key=f"dl_pdf_current_{sop_fingerprint(current_sop)}",
        )


# --- Step 2: Quality pass (Review & Fix) ---
last_sop = st.session_state.get("last_sop_text", "")
if api_key and last_sop:
    st.divider()
    st.subheader("Review & Fix SOP (Quality pass)")
    st.caption("Runs an editor pass to fix gaps, unclear steps, and missing roles/records.")

    do_review = st.button("Review & Fix SOP", type="secondary")
    if do_review:
        with st.spinner("Reviewing and improving SOP..."):
            try:
                fixed = review_and_fix_sop_cached(
                    api_key=api_key,
                    model=model,
                    temperature=min(float(temperature), 0.4),
                    sop_text=last_sop,
                    strictness=strictness,
                    tone=tone,
                    compliance_standard=compliance_standard.strip(),
                )
                st.session_state.last_fixed_sop_text = fixed

                inferred_topic = st.session_state.get("last_inferred_topic", "SOP")
                add_to_history(
                    {
                        "ts": datetime.now(timezone.utc).isoformat(),
                        "label": f"{inferred_topic} (revised)",
                        "template_name": template_name,
                        "source": "revised",
                        "sop_text": fixed,
                        "sop_sha256": sop_fingerprint(fixed),
                    }
                )
            except Exception as e:
                show_busy_error()

    fixed_sop = st.session_state.get("last_fixed_sop_text", "")
    if fixed_sop:
        st.subheader("Revised SOP")
        st.markdown(fixed_sop)

        # Make revised SOP the current editable SOP by default.
        st.session_state.current_sop_text = fixed_sop

        inferred_topic = st.session_state.get("last_inferred_topic", "SOP")
        safe_name = (
            "".join(c for c in str(inferred_topic).strip() if c.isalnum() or c in (" ", "-", "_")).strip()
            or "sop"
        )

        with st.expander("Interactive Step Editor (edit revised SOP)", expanded=False):
            st.caption("Edit the revised SOP here, then download the edited version.")
            edited_rev = st.text_area(
                "Edit revised SOP text",
                value=st.session_state.current_sop_text,
                height=320,
                key=f"editor_rev_{sop_fingerprint(fixed_sop)}",
            )
            col_r1, col_r2 = st.columns(2)
            with col_r1:
                if st.button("Save revised edits", key=f"save_rev_{sop_fingerprint(fixed_sop)}"):
                    st.session_state.current_sop_text = edited_rev
                    st.success("Edits saved. Revised downloads will use the edited SOP.")
            with col_r2:
                if st.button("Reset to revised", key=f"reset_rev_{sop_fingerprint(fixed_sop)}"):
                    st.session_state.current_sop_text = fixed_sop
                    st.success("Reset to the revised SOP.")

        sop_for_download = st.session_state.get("current_sop_text") or fixed_sop
        try:
            pdf_bytes = create_pdf_bytes(sop_for_download)
            docx_bytes = create_docx_bytes(safe_name, sop_for_download)
        except Exception:
            pdf_bytes = b""
            docx_bytes = b""
            show_busy_error()

        col_c, col_d = st.columns(2)
        with col_c:
            if pdf_bytes:
                st.download_button(
                    "Download Revised PDF",
                    data=pdf_bytes,
                    file_name=f"{safe_name}-revised.pdf",
                    mime="application/pdf",
                    key=f"dl_pdf_rev_{sop_fingerprint(sop_for_download)}",
                )
        with col_d:
            if docx_bytes:
                st.download_button(
                    "Download Revised DOCX",
                    data=docx_bytes,
                    file_name=f"{safe_name}-revised.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"dl_docx_rev_{sop_fingerprint(sop_for_download)}",
                )

        st.markdown("### Rate the revised SOP")
        rating2 = st.radio(
            "Was the revised SOP helpful?",
            ["👍 Thumbs Up", "👎 Thumbs Down"],
            horizontal=True,
            key=f"rating_rev_{sop_fingerprint(fixed_sop)}",
        )
        reason2 = ""
        if rating2.startswith("👎"):
            reason2 = st.text_area(
                "What should still be improved?",
                placeholder="Be specific about what’s missing or unclear.",
                key=f"reason_rev_{sop_fingerprint(fixed_sop)}",
            )
        if st.button("Submit revised feedback", key=f"submit_rev_{sop_fingerprint(fixed_sop)}"):
            append_feedback(
                {
                    "ts": datetime.now(timezone.utc).isoformat(),
                    "rating": "up" if rating2.startswith("👍") else "down",
                    "reason": reason2.strip(),
                    "template_name": template_name,
                    "strictness": strictness,
                    "tone": tone,
                    "compliance_standard": compliance_standard or "",
                    "audience": (audience or "").strip(),
                    "tools_used": (tools_used or "").strip(),
                    "include_definitions": bool(include_definitions),
                    "include_safety_compliance": bool(include_safety_compliance),
                    "include_records": bool(include_records),
                    "include_checklist": bool(include_checklist),
                    "model": model,
                    "temperature": float(temperature),
                    "notes_chars": len((st.session_state.get("notes") or "").strip()),
                    "sop_sha256": sop_fingerprint(fixed_sop),
                    "source": "revised",
                }
            )
            st.success("Thanks — feedback saved.")


# --- Compliance Auditor (AI Critic) ---
candidate_sop_for_audit = st.session_state.get("last_fixed_sop_text") or st.session_state.get("last_sop_text") or ""
if api_key and candidate_sop_for_audit:
    st.divider()
    st.subheader("Compliance Auditor (AI Critic)")
    st.caption("A second AI pass that critiques the SOP for gaps, risks, and missing compliance items.")

    run_audit = st.button("Review SOP (Auditor)")
    if run_audit:
        with st.spinner("Auditing SOP..."):
            try:
                audit = compliance_audit_cached(
                    api_key=api_key,
                    model=model,
                    temperature=0.2,
                    sop_text=candidate_sop_for_audit,
                    template_name=template_name,
                    strictness=strictness,
                    tone=tone,
                    compliance_standard=compliance_standard.strip(),
                    company_rules_context=st.session_state.get("company_rules_context", "").strip(),
                )
                st.session_state.last_audit_text = audit
            except Exception as e:
                show_busy_error()

    audit_text = st.session_state.get("last_audit_text", "")
    if audit_text:
        st.markdown(audit_text)


# --- Visual Flowchart ---
candidate_sop_for_flowchart = st.session_state.get("last_fixed_sop_text") or st.session_state.get("last_sop_text") or ""
if api_key and candidate_sop_for_flowchart:
    st.divider()
    st.subheader("Visual flowchart")
    st.caption("Generates a flowchart from the latest SOP (revised if available).")

    gen_chart = st.button("Generate Flowchart")
    if gen_chart:
        with st.spinner("Generating flowchart..."):
            try:
                mermaid_code = generate_flowchart_mermaid_cached(
                    api_key=api_key,
                    model=model,
                    temperature=0.2,
                    sop_text=candidate_sop_for_flowchart,
                )
                st.session_state.last_mermaid_flowchart = mermaid_code
            except Exception as e:
                show_busy_error()

    mermaid_code = st.session_state.get("last_mermaid_flowchart", "")
    if mermaid_code:
        render_mermaid(mermaid_code, height_px=700)
        st.download_button(
            "Download Flowchart (Mermaid)",
            data=mermaid_code.encode("utf-8"),
            file_name="sop-flowchart.mmd",
            mime="text/plain",
        )
